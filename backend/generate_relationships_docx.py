"""
Generate professional Word document explaining all FederCare database
table relationships.
Output: federcare_database_relationships.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DARK_BLUE = RGBColor(26, 60, 110)
HEADER_FILL = 'C5D8F6'
ALT_ROW_FILL = 'F2F7FE'


doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Header
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header_para.add_run('FederCare: AI Health Network')
header_run.italic = True
header_run.font.size = Pt(10)
header_run.font.color.rgb = DARK_BLUE

# Footer
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
footer_run = footer_para.add_run(
    'Mar Thoma Institute of Information Technology, Ayur'
)
footer_run.italic = True
footer_run.bold = True
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = DARK_BLUE

footer_para.add_run('\t\t\t\t\t\t')

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

page_run = footer_para.add_run()
page_run.element.append(fldChar1)
page_run.element.append(instrText)
page_run.element.append(fldChar2)
page_run.font.size = Pt(9)
page_run.bold = True


def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_BLUE


def add_explanation(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.size = Pt(10)


def add_steps(doc, steps):
    for step in steps:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(step)
        run.font.size = Pt(10)


def add_relationship_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        set_cell_background(header_cells[i], HEADER_FILL)
        set_cell_borders(header_cells[i])
        para = header_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        for j, val in enumerate(row):
            row_cells[j].text = val
            bg = 'FFFFFF' if idx % 2 == 0 else ALT_ROW_FILL
            set_cell_background(row_cells[j], bg)
            set_cell_borders(row_cells[j])
            para = row_cells[j].paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(val)
            run.font.size = Pt(9)
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# ===== TITLE =====
doc.add_paragraph('\n\n')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('Database Relationships and Entity Connections')
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = DARK_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('FederCare: AI Health Network')
sr.bold = True
sr.font.size = Pt(14)

doc.add_paragraph()

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dept.add_run('Department of Computer Applications')
dr.font.size = Pt(12)

college = doc.add_paragraph()
college.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = college.add_run(
    'Mar Thoma Institute of Information Technology, Ayur'
)
cr.font.size = Pt(12)

year = doc.add_paragraph()
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
yr = year.add_run('2025 - 2026')
yr.font.size = Pt(12)

doc.add_page_break()


# ===== SECTION 1 =====
add_section_heading(doc, '1. Database Overview')
add_explanation(doc,
    'FederCare: AI Health Network uses a relational database with 41 '
    'tables organized across 10 Django applications. All tables use '
    'UUID as primary keys for security and scalability. The central '
    'table is login_credentials which connects all user roles in the '
    'system.'
)


# ===== SECTION 2 =====
add_section_heading(doc, '2. Core Authentication Layer')
add_explanation(doc,
    'The authentication layer is the foundation of the entire system. '
    'Every user regardless of their role must have a login_credentials '
    'record. This single table connects to all 8 user profile tables '
    'through a One-to-One relationship.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Purpose'],
    [
        ['login_credentials', 'One-to-One', 'super_admin', 'Super admin profile'],
        ['login_credentials', 'One-to-One', 'hospital_registrations', 'Hospital profile'],
        ['login_credentials', 'One-to-One', 'doctor_registrations', 'Doctor profile'],
        ['login_credentials', 'One-to-One', 'patient_registrations', 'Patient profile'],
        ['login_credentials', 'One-to-One', 'pharmacist_registrations', 'Pharmacist profile'],
        ['login_credentials', 'One-to-One', 'lab_tech_registrations', 'Lab tech profile'],
        ['login_credentials', 'One-to-One', 'ambulance_driver_registrations', 'Driver profile'],
        ['login_credentials', 'One-to-One', 'vendor_registrations', 'Vendor profile'],
    ]
)
add_explanation(doc,
    'When a user logs in, the system checks login_credentials to verify '
    'identity. Based on the role field, it loads the corresponding '
    'profile table. For example, a doctor login loads doctor_registrations '
    'to get specialization, hospital, and consultation fee details.'
)


# ===== SECTION 3 =====
add_section_heading(doc, '3. Hospital Module Relationships')
add_explanation(doc,
    'The hospital is the central entity that connects doctors, lab '
    'technicians, drivers, beds, departments, and inventory. Everything '
    'in FederCare belongs to or is associated with a hospital.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['hospital_registrations', 'One-to-Many', 'departments', 'Each hospital has multiple departments'],
        ['hospital_registrations', 'One-to-Many', 'beds', 'Each hospital has multiple beds'],
        ['hospital_registrations', 'One-to-Many', 'hospital_inventory', 'Each hospital has inventory items'],
        ['hospital_registrations', 'One-to-Many', 'hospital_patients', 'Training data for FL per hospital'],
        ['hospital_registrations', 'One-to-Many', 'doctor_registrations', 'Doctors belong to a hospital'],
        ['hospital_registrations', 'One-to-Many', 'lab_tech_registrations', 'Lab techs belong to a hospital'],
        ['hospital_registrations', 'One-to-Many', 'ambulance_driver_registrations', 'Drivers belong to a hospital'],
        ['hospital_registrations', 'One-to-Many', 'ambulances', 'Ambulances registered to hospital'],
        ['hospital_registrations', 'One-to-Many', 'fl_hospital_weights', 'FL weights submitted per hospital'],
    ]
)
add_explanation(doc,
    'When a hospital admin logs in, they can manage all resources linked '
    'to their hospital_id. For example, adding a bed creates a record in '
    'beds table with hospital_id pointing to their hospital. Doctors can '
    'only see patients who had consultations in their hospital.'
)


# ===== SECTION 4 =====
add_section_heading(doc, '4. Patient Journey Relationships')
add_explanation(doc,
    'The patient_registrations table is the most connected table in the '
    'system. A patient can book consultations, order medicines, book lab '
    'tests, upload medical images, trigger emergencies, and receive '
    'prescriptions. All these actions create records linked to patient_id.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['patient_registrations', 'One-to-Many', 'consultations', 'Patient books multiple consultations'],
        ['patient_registrations', 'One-to-Many', 'prescriptions', 'Patient receives prescriptions'],
        ['patient_registrations', 'One-to-Many', 'lab_test_orders', 'Patient books lab tests'],
        ['patient_registrations', 'One-to-Many', 'medicine_orders', 'Patient orders medicines'],
        ['patient_registrations', 'One-to-Many', 'ehr_records', 'Patient EHR records'],
        ['patient_registrations', 'One-to-Many', 'ehr_images', 'Patient medical images'],
        ['patient_registrations', 'One-to-Many', 'allergies', 'Patient allergy records'],
        ['patient_registrations', 'One-to-Many', 'risk_assessments', 'AI risk assessments'],
        ['patient_registrations', 'One-to-Many', 'ehr_consent_log', 'EHR access consents'],
        ['patient_registrations', 'One-to-Many', 'emergency_requests', 'Emergency SOS requests'],
        ['patient_registrations', 'One-to-Many', 'patient_complaints', 'Filed complaints'],
        ['patient_registrations', 'One-to-Many', 'triage_sessions', 'AI symptom triage'],
    ]
)
add_explanation(doc, 'A complete patient journey works like this:')
add_steps(doc, [
    'Patient registers — patient_registrations created.',
    'Patient books consultation — consultations record created linked to patient_id, doctor_id, and slot_id.',
    'Doctor issues prescription — prescriptions record created linked to consultation_id and patient_id.',
    'Patient orders medicine — medicine_orders created linked to patient_id and prescription_id.',
    'Patient books lab test — lab_test_orders created.',
    'Results uploaded — ehr_records updated.',
    'Emergency triggered — emergency_requests created.',
])


# ===== SECTION 5 =====
add_section_heading(doc, '5. Doctor Consultation Flow')
add_explanation(doc,
    'The consultation system links doctors, patients, time slots, '
    'prescriptions, and lab orders in a chain.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['doctor_registrations', 'One-to-Many', 'doctor_slots', 'Doctor creates available time slots'],
        ['doctor_registrations', 'One-to-Many', 'consultations', 'Doctor has many consultations'],
        ['doctor_registrations', 'One-to-Many', 'prescriptions', 'Doctor issues prescriptions'],
        ['doctor_registrations', 'One-to-Many', 'lab_orders', 'Doctor creates lab orders'],
        ['doctor_slots', 'One-to-One', 'consultations', 'One slot per consultation booking'],
        ['consultations', 'One-to-Many', 'prescriptions', 'Consultation can have prescriptions'],
        ['consultations', 'One-to-Many', 'lab_orders', 'Consultation can have lab orders'],
    ]
)
add_explanation(doc, 'Step by step consultation flow:')
add_steps(doc, [
    'Doctor creates slots in doctor_slots table with slot_date, start_time, end_time.',
    'Patient selects a slot and books — consultation created linking patient_id + doctor_id + slot_id, and slot is_booked set to True.',
    'During video consultation doctor writes notes saved to consultations table.',
    'Doctor issues prescription — prescriptions record created linked to consultation_id.',
    'Doctor creates lab order — lab_orders created linked to consultation_id.',
])


# ===== SECTION 6 =====
add_section_heading(doc, '6. Emergency SOS Chain')
add_explanation(doc,
    'The emergency system follows a clear chain from patient SOS to '
    'ambulance dispatch and hospital delivery.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['emergency_requests', 'Many-to-One', 'patient_registrations', 'Patient who triggered SOS'],
        ['emergency_requests', 'Many-to-One', 'hospital_registrations', 'Nearest hospital assigned'],
        ['emergency_requests', 'Many-to-One', 'beds', 'Bed reserved for patient'],
        ['emergency_requests', 'One-to-Many', 'ambulance_dispatch', 'Dispatches for emergency'],
        ['ambulance_dispatch', 'Many-to-One', 'ambulances', 'Ambulance assigned'],
        ['ambulances', 'Many-to-One', 'ambulance_driver_registrations', 'Driver of ambulance'],
    ]
)
add_explanation(doc, 'Emergency chain explanation:')
add_steps(doc, [
    'Patient triggers SOS — emergency_requests record created with patient_lat, patient_lng, severity.',
    'System finds nearest ambulance using GPS coordinates in ambulances table — distance calculated using Haversine formula.',
    'AmbulanceDispatch record created linking emergency_id to ambulance_id.',
    'Ambulance.is_available set to False — prevents double booking.',
    'Driver notified via WebSocket.',
    'Driver updates status: dispatched → en_route → arrived → completed.',
    'Hospital Admin acknowledges patient — emergency status = completed, ambulance.is_available = True again.',
])


# ===== SECTION 7 =====
add_section_heading(doc, '7. Federated Learning Chain')
add_explanation(doc,
    'FederCare uses Federated Learning (FL) to train AI models across '
    'hospitals without sharing patient data. The FL system has its own '
    'chain of tables.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['fl_global_models', 'One-to-Many', 'fl_rounds', 'Global model has multiple rounds'],
        ['fl_rounds', 'One-to-Many', 'fl_hospital_weights', 'Each round collects hospital weights'],
        ['fl_hospital_weights', 'Many-to-One', 'hospital_registrations', 'Weight submitted by hospital'],
        ['hospital_patients', 'Many-to-One', 'hospital_registrations', 'Training data per hospital'],
        ['triage_sessions', 'Many-to-One', 'patient_registrations', 'Patient symptom triage'],
    ]
)
add_explanation(doc, 'FL training flow:')
add_steps(doc, [
    'Super Admin initializes FL round — fl_global_models record created and fl_rounds record created with status=active.',
    'Each hospital trains local model using their hospital_patients records (real patient data in hospital).',
    'Hospital submits weights — fl_hospital_weights record created linking round_id and hospital_id.',
    'When threshold hospitals submit — FedAvg algorithm runs and weights averaged to create new global model.',
    'New fl_global_models record created with updated accuracy.',
    'All hospitals use updated model for AI symptom checking.',
])


# ===== SECTION 8 =====
add_section_heading(doc, '8. Medicine Order Chain')
add_explanation(doc,
    'The medicine ordering system connects patients, pharmacists, '
    'prescriptions, and inventory in a verified flow.'
)
add_relationship_table(
    doc,
    ['Table', 'Relationship', 'Connected To', 'Description'],
    [
        ['medicine_orders', 'Many-to-One', 'patient_registrations', 'Patient who ordered'],
        ['medicine_orders', 'Many-to-One', 'pharmacist_registrations', 'Pharmacist handling order'],
        ['medicine_orders', 'Many-to-One', 'prescriptions', 'Linked prescription if required'],
        ['pharmacy_inventory', 'Many-to-One', 'pharmacist_registrations', 'Medicines per pharmacy'],
    ]
)
add_explanation(doc, 'Medicine order flow:')
add_steps(doc, [
    'Patient browses pharmacy_inventory — medicines linked to pharmacist_id.',
    'Patient adds medicines to cart — system checks requires_prescription field in pharmacy_inventory.',
    'If prescription needed: patient uploads prescription file — medicine_orders status = prescription_uploaded.',
    'Pharmacist verifies prescription — medicine_orders.prescription_verified = True and payment_enabled = True.',
    'Patient pays via Razorpay — payment_status = paid and order_status = confirmed.',
    'Pharmacist dispatches — delivery_otp generated and dispatched_at timestamp set.',
    'Patient enters OTP — otp_verified = True and delivered_at timestamp set.',
])


# ===== SECTION 9 =====
add_section_heading(doc, '9. Key Design Decisions')
add_relationship_table(
    doc,
    ['Design Decision', 'Reason'],
    [
        ['UUID Primary Keys', 'More secure than integers. Prevents ID enumeration attacks. Works across distributed systems.'],
        ['Single Login Table', 'All 8 user types share one login table. Simplifies authentication. Role-based access from one place.'],
        ['JSON Fields for Lists', 'Symptoms, medicines, test results stored as JSON. Flexible schema for varying data structures.'],
        ['Soft Delete Pattern', 'Records use status fields instead of deletion. Maintains audit trail and data integrity.'],
        ['OTP Delivery System', 'Equipment and medicine deliveries use OTP verification. Prevents fake delivery confirmations.'],
        ['Razorpay Integration', 'Payment IDs stored in order tables. Links payment gateway records to system orders.'],
        ['WebSocket Real-time', 'Notifications and GPS updates use Django Channels WebSocket for instant updates.'],
        ['Differential Privacy', 'FL weights include noise_added flag. Protects individual patient data during training.'],
        ['Audit Logging', 'Every action logged in audit_logs. Full traceability for compliance and debugging.'],
        ['Consent-based EHR', 'Doctor accesses EHR only with patient QR consent. ehr_consent_log tracks all access.'],
    ]
)


# ===== SECTION 10 =====
add_section_heading(doc, '10. Complete Relationship Summary')
add_relationship_table(
    doc,
    ['Module', 'Tables', 'Key Relationships'],
    [
        ['Auth Module', 'login_credentials, login_sessions, audit_logs, notifications, role_permissions', 'login_credentials is central hub for all 8 user types'],
        ['Hospital Module', 'hospital_registrations, departments, beds, hospital_inventory, hospital_patients', 'hospital_id links all hospital resources'],
        ['Patient Module', 'patient_registrations, ehr_records, ehr_images, allergies, risk_assessments, ehr_consent_log', 'patient_id links all health records'],
        ['Consultation Module', 'doctor_slots, consultations, prescriptions, lab_orders, lab_reports', 'slot → consultation → prescription chain'],
        ['Pharmacy Module', 'pharmacist_registrations, medicine_orders, pharmacy_inventory', 'prescription verification before payment'],
        ['Lab Module', 'lab_tech_registrations, lab_orders, lab_reports, lab_test_orders', 'doctor orders → lab tech uploads → patient downloads'],
        ['Emergency Module', 'emergency_requests, ambulance_dispatch, ambulances, ambulance_driver_registrations', 'GPS-based nearest ambulance dispatch'],
        ['Vendor Module', 'vendor_registrations, equipment_catalog, equipment_orders', 'OTP-based equipment delivery'],
        ['FL Module', 'fl_global_models, fl_rounds, fl_hospital_weights, triage_sessions, epidemic_trends', 'FedAvg aggregation from hospital weights'],
        ['AI Module', 'triage_sessions, risk_assessments, epidemic_trends', 'ML models trained on hospital_patients data'],
    ]
)


output_path = 'federcare_database_relationships.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
print('Sections: 10')
