"""
Generate professional Word document with all FederCare database tables.
Output: federcare_database_tables.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Header
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header_para.add_run('FederCare: AI Health Network')
header_run.italic = True
header_run.font.size = Pt(10)
header_run.font.color.rgb = RGBColor(26, 60, 110)

# Footer
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
footer_run = footer_para.add_run(
    'Mar Thoma Institute of Information Technology, Ayur'
)
footer_run.italic = True
footer_run.bold = True
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(26, 60, 110)

footer_para.add_run('\t\t\t\t\t\t')

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

page_run = footer_para.add_run()
page_run.element.append(fldChar1)
page_run.element.append(instrText)
page_run.element.append(fldChar2)
page_run.font.size = Pt(9)
page_run.bold = True


def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_table_to_doc(doc, table_name, description, fields):
    heading_para = doc.add_paragraph()
    heading_para.paragraph_format.space_before = Pt(12)
    heading_para.paragraph_format.space_after = Pt(2)
    heading_run = heading_para.add_run(f'Table Name: {table_name}')
    heading_run.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    desc_para = doc.add_paragraph()
    desc_para.paragraph_format.space_before = Pt(0)
    desc_para.paragraph_format.space_after = Pt(6)
    desc_run = desc_para.add_run(f'Description: {description}')
    desc_run.bold = True
    desc_run.italic = True
    desc_run.font.size = Pt(10)

    col_widths = [Cm(3.5), Cm(3.0), Cm(4.0), Cm(6.0)]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    header_cells = table.rows[0].cells
    headers = ['Attribute', 'Data Type', 'Constraints', 'Description']

    for i, (cell, header) in enumerate(zip(header_cells, headers)):
        cell.text = header
        set_cell_background(cell, 'C5D8F6')
        set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, field in enumerate(fields):
        row_cells = table.add_row().cells
        row_cells[0].text = field[0]
        row_cells[1].text = field[1]
        row_cells[2].text = field[2]
        row_cells[3].text = field[3]

        bg_color = 'FFFFFF' if idx % 2 == 0 else 'F2F7FE'

        for cell in row_cells:
            set_cell_background(cell, bg_color)
            set_cell_borders(cell)
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


# Title page
doc.add_paragraph('\n\n\n\n')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('FederCare: AI Health Network')
title_run.bold = True
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(26, 60, 110)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run('Database Design Document')
sub_run.bold = True
sub_run.font.size = Pt(14)

doc.add_paragraph()

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
dept_run = dept.add_run('Department of Computer Applications')
dept_run.font.size = Pt(12)

college = doc.add_paragraph()
college.alignment = WD_ALIGN_PARAGRAPH.CENTER
college_run = college.add_run(
    'Mar Thoma Institute of Information Technology, Ayur'
)
college_run.font.size = Pt(12)

year = doc.add_paragraph()
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
year_run = year.add_run('2025 - 2026')
year_run.font.size = Pt(12)

doc.add_page_break()


tables_data = [
    {
        'name': 'login_credentials',
        'description': 'Stores login credentials and authentication details for all system users.',
        'fields': [
            ('login_id', 'UUID', 'Primary Key, Not Null', 'Unique login identifier'),
            ('email', 'Varchar(100)', 'Unique, Not Null', 'User email address'),
            ('password_hash', 'Varchar(255)', 'Not Null', 'Encrypted password hash'),
            ('role', 'Varchar(40)', 'Not Null', 'User role (super_admin/doctor/patient etc.)'),
            ('is_active', 'Boolean', 'Not Null', 'Account active status'),
            ('is_approved', 'Boolean', 'Not Null', 'Account approval status'),
            ('last_login', 'DateTime', 'Null', 'Last login timestamp'),
            ('login_attempts', 'Integer', 'Not Null', 'Failed login attempts count'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'super_admin',
        'description': 'Stores super administrator profile details.',
        'fields': [
            ('admin_id', 'UUID', 'Primary Key, Not Null', 'Unique admin identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Administrator full name'),
            ('phone', 'Varchar(15)', 'Not Null', 'Contact phone number'),
            ('profile_photo', 'Varchar(500)', 'Not Null', 'Profile photo URL'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'role_permissions',
        'description': 'Stores module-level permissions for each user role.',
        'fields': [
            ('permission_id', 'UUID', 'Primary Key, Not Null', 'Unique permission identifier'),
            ('role', 'Varchar(40)', 'Not Null', 'User role name'),
            ('module', 'Varchar(50)', 'Not Null', 'System module name'),
            ('can_read', 'Boolean', 'Not Null', 'Read permission flag'),
            ('can_write', 'Boolean', 'Not Null', 'Write permission flag'),
            ('can_delete', 'Boolean', 'Not Null', 'Delete permission flag'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'login_sessions',
        'description': 'Tracks active login sessions and JWT tokens for security.',
        'fields': [
            ('session_id', 'UUID', 'Primary Key, Not Null', 'Unique session identifier'),
            ('login_id', 'UUID', 'FK to login_credentials', 'Reference to user login'),
            ('jwt_token_hash', 'Varchar(255)', 'Not Null', 'Hashed JWT token'),
            ('device_info', 'Text', 'Not Null', 'Client device information'),
            ('ip_address', 'GenericIP', 'Null', 'Client IP address'),
            ('expires_at', 'DateTime', 'Not Null', 'Session expiry timestamp'),
            ('created_at', 'DateTime', 'Not Null', 'Session creation timestamp'),
        ]
    },
    {
        'name': 'audit_logs',
        'description': 'Records all system actions for security and compliance auditing.',
        'fields': [
            ('log_id', 'UUID', 'Primary Key, Not Null', 'Unique log identifier'),
            ('login_id', 'UUID', 'FK to login_credentials, Null', 'User who performed action'),
            ('action', 'Varchar(255)', 'Not Null', 'Action performed'),
            ('module', 'Varchar(50)', 'Not Null', 'Module where action occurred'),
            ('entity_type', 'Varchar(50)', 'Not Null', 'Type of entity affected'),
            ('entity_id', 'UUID', 'Null', 'ID of affected entity'),
            ('old_value', 'JSON', 'Null', 'Previous value before change'),
            ('new_value', 'JSON', 'Null', 'New value after change'),
            ('ip_address', 'GenericIP', 'Null', 'IP address of request'),
            ('logged_at', 'DateTime', 'Not Null', 'Action timestamp'),
        ]
    },
    {
        'name': 'notifications',
        'description': 'Stores system notifications sent to users across all roles.',
        'fields': [
            ('notif_id', 'UUID', 'Primary Key, Not Null', 'Unique notification identifier'),
            ('login_id', 'UUID', 'FK to login_credentials', 'Recipient user reference'),
            ('title', 'Varchar(200)', 'Not Null', 'Notification title'),
            ('message', 'Text', 'Not Null', 'Notification message content'),
            ('notif_type', 'Varchar(30)', 'Not Null', 'Notification type'),
            ('is_read', 'Boolean', 'Not Null', 'Read status flag'),
            ('related_id', 'UUID', 'Null', 'Related entity identifier'),
            ('created_at', 'DateTime', 'Not Null', 'Notification creation timestamp'),
        ]
    },
    {
        'name': 'hospital_registrations',
        'description': 'Stores hospital registration and profile information.',
        'fields': [
            ('hospital_id', 'UUID', 'Primary Key, Not Null', 'Unique hospital identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('hospital_name', 'Varchar(200)', 'Not Null', 'Official hospital name'),
            ('registration_no', 'Varchar(100)', 'Unique, Not Null', 'Government registration number'),
            ('address', 'Text', 'Not Null', 'Hospital address'),
            ('city', 'Varchar(100)', 'Not Null', 'City name'),
            ('state', 'Varchar(100)', 'Not Null', 'State name'),
            ('latitude', 'Decimal(9,7)', 'Null', 'GPS latitude coordinate'),
            ('longitude', 'Decimal(9,7)', 'Null', 'GPS longitude coordinate'),
            ('contact_phone', 'Varchar(15)', 'Not Null', 'Hospital contact number'),
            ('contact_email', 'Varchar(100)', 'Not Null', 'Hospital contact email'),
            ('telemedicine_enabled', 'Boolean', 'Not Null', 'Telemedicine support flag'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'departments',
        'description': 'Stores hospital department information.',
        'fields': [
            ('dept_id', 'UUID', 'Primary Key, Not Null', 'Unique department identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Parent hospital reference'),
            ('dept_name', 'Varchar(100)', 'Not Null', 'Department name'),
            ('description', 'Text', 'Not Null', 'Department description'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'beds',
        'description': 'Stores hospital bed information and availability status.',
        'fields': [
            ('bed_id', 'UUID', 'Primary Key, Not Null', 'Unique bed identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Parent hospital reference'),
            ('bed_type', 'Varchar(50)', 'Not Null', 'Bed type (general/ICU/emergency)'),
            ('ward_name', 'Varchar(100)', 'Not Null', 'Ward name'),
            ('status', 'Varchar(20)', 'Not Null', 'Bed status (available/occupied/reserved)'),
            ('reserved_for', 'UUID', 'FK to patient_registrations, Null', 'Reserved patient reference'),
            ('reserved_at', 'DateTime', 'Null', 'Reservation timestamp'),
            ('admitted_at', 'DateTime', 'Null', 'Patient admission timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'hospital_inventory',
        'description': 'Tracks medical equipment and supply inventory for hospitals.',
        'fields': [
            ('inventory_id', 'UUID', 'Primary Key, Not Null', 'Unique inventory identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Parent hospital reference'),
            ('item_name', 'Varchar(200)', 'Not Null', 'Item name'),
            ('category', 'Varchar(50)', 'Not Null', 'Item category'),
            ('quantity', 'Integer', 'Not Null', 'Available quantity'),
            ('unit', 'Varchar(20)', 'Not Null', 'Unit of measurement'),
            ('reorder_level', 'Integer', 'Not Null', 'Minimum reorder threshold'),
            ('last_restocked', 'DateTime', 'Null', 'Last restock timestamp'),
            ('maintenance_due', 'Date', 'Null', 'Next maintenance date'),
        ]
    },
    {
        'name': 'hospital_patients',
        'description': 'Stores patient records added by hospitals for federated learning training data.',
        'fields': [
            ('patient_id', 'UUID', 'Primary Key, Not Null', 'Unique record identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Parent hospital reference'),
            ('added_by', 'UUID', 'FK to login_credentials, Null', 'Staff who added record'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Patient full name'),
            ('age', 'Integer', 'Not Null', 'Patient age'),
            ('gender', 'Varchar(10)', 'Not Null', 'Patient gender'),
            ('blood_group', 'Varchar(5)', 'Not Null', 'Patient blood group'),
            ('symptoms', 'JSON', 'Not Null', 'Symptoms list'),
            ('diagnosis', 'Varchar(200)', 'Not Null', 'Medical diagnosis'),
            ('visit_date', 'Date', 'Not Null', 'Hospital visit date'),
            ('notes', 'Text', 'Not Null', 'Clinical notes'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'patient_registrations',
        'description': 'Stores registered patient profile and health information.',
        'fields': [
            ('patient_id', 'UUID', 'Primary Key, Not Null', 'Unique patient identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Patient full name'),
            ('dob', 'Date', 'Not Null', 'Date of birth'),
            ('gender', 'Varchar(10)', 'Not Null', 'Patient gender'),
            ('blood_group', 'Varchar(5)', 'Not Null', 'Blood group'),
            ('height_cm', 'Decimal', 'Null', 'Height in centimeters'),
            ('weight_kg', 'Decimal', 'Null', 'Weight in kilograms'),
            ('bmi', 'Decimal', 'Null', 'Body Mass Index'),
            ('address', 'Text', 'Not Null', 'Residential address'),
            ('emergency_contact', 'Varchar(15)', 'Not Null', 'Emergency contact number'),
            ('qr_code_url', 'Varchar(500)', 'Not Null', 'EHR QR code URL'),
            ('lifestyle_data', 'JSON', 'Not Null', 'Lifestyle and health habits data'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'ehr_records',
        'description': 'Stores Electronic Health Records for patients.',
        'fields': [
            ('record_id', 'UUID', 'Primary Key, Not Null', 'Unique record identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('added_by', 'UUID', 'FK to login_credentials, Null', 'Staff who added record'),
            ('record_type', 'Varchar(50)', 'Not Null', 'Record type (prescription/lab/diagnosis)'),
            ('title', 'Varchar(200)', 'Not Null', 'Record title'),
            ('content', 'Text', 'Not Null', 'Record content'),
            ('file_url', 'Varchar(500)', 'Not Null', 'Attached file URL'),
            ('is_sensitive', 'Boolean', 'Not Null', 'Sensitive data flag'),
            ('recorded_at', 'DateTime', 'Not Null', 'Record timestamp'),
        ]
    },
    {
        'name': 'allergies',
        'description': 'Stores patient allergy information for medical safety.',
        'fields': [
            ('allergy_id', 'UUID', 'Primary Key, Not Null', 'Unique allergy identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('allergen', 'Varchar(100)', 'Not Null', 'Allergen name'),
            ('reaction', 'Text', 'Not Null', 'Reaction description'),
            ('severity', 'Varchar(20)', 'Not Null', 'Severity (mild/moderate/severe)'),
            ('noted_by', 'UUID', 'FK to doctor_registrations, Null', 'Doctor who noted allergy'),
            ('noted_at', 'DateTime', 'Not Null', 'Date allergy was noted'),
        ]
    },
    {
        'name': 'ehr_consent_log',
        'description': 'Tracks patient consent for EHR data access by doctors.',
        'fields': [
            ('consent_id', 'UUID', 'Primary Key, Not Null', 'Unique consent identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('accessed_by', 'UUID', 'FK to login_credentials', 'Staff who accessed EHR'),
            ('access_type', 'Varchar(50)', 'Not Null', 'Access type (view/download)'),
            ('data_shared', 'JSON', 'Not Null', 'List of shared data fields'),
            ('consent_given', 'Boolean', 'Not Null', 'Consent status flag'),
            ('expires_at', 'DateTime', 'Null', 'Consent expiry timestamp'),
            ('accessed_at', 'DateTime', 'Not Null', 'Access timestamp'),
        ]
    },
    {
        'name': 'risk_assessments',
        'description': 'Stores AI-generated health risk assessment results for patients.',
        'fields': [
            ('risk_id', 'UUID', 'Primary Key, Not Null', 'Unique assessment identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('diabetes_risk', 'Decimal', 'Null', 'Diabetes risk percentage'),
            ('heart_risk', 'Decimal', 'Null', 'Heart disease risk percentage'),
            ('hypertension_risk', 'Decimal', 'Null', 'Hypertension risk percentage'),
            ('risk_level', 'Varchar(20)', 'Not Null', 'Overall risk level (low/medium/high)'),
            ('recommendations', 'Text', 'Not Null', 'AI recommendations'),
            ('alert_sent', 'Boolean', 'Not Null', 'Alert notification sent flag'),
            ('assessed_at', 'DateTime', 'Not Null', 'Assessment timestamp'),
        ]
    },
    {
        'name': 'patient_complaints',
        'description': 'Stores patient complaints about doctors, hospitals or vendors.',
        'fields': [
            ('complaint_id', 'UUID', 'Primary Key, Not Null', 'Unique complaint identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Complainant patient reference'),
            ('complaint_type', 'Varchar(30)', 'Not Null', 'Complaint type (doctor/vendor/hospital)'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations, Null', 'Complained doctor reference'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations, Null', 'Complained hospital reference'),
            ('vendor_id', 'UUID', 'FK to vendor_registrations, Null', 'Complained vendor reference'),
            ('subject', 'Varchar(200)', 'Not Null', 'Complaint subject'),
            ('description', 'Text', 'Not Null', 'Detailed complaint description'),
            ('status', 'Varchar(20)', 'Not Null', 'Status (pending/reviewed/resolved)'),
            ('admin_response', 'Text', 'Not Null', 'Admin response to complaint'),
            ('created_at', 'DateTime', 'Not Null', 'Complaint submission timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'lab_test_orders',
        'description': 'Stores patient lab test bookings and results.',
        'fields': [
            ('order_id', 'UUID', 'Primary Key, Not Null', 'Unique order identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations, Null', 'Hospital reference'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations, Null', 'Referring doctor reference'),
            ('tests', 'JSON', 'Not Null', 'List of ordered tests'),
            ('total_fee', 'Decimal', 'Not Null', 'Total test fee'),
            ('appointment_date', 'Date', 'Null', 'Scheduled appointment date'),
            ('appointment_time', 'Time', 'Null', 'Scheduled appointment time'),
            ('status', 'Varchar(20)', 'Not Null', 'Order status'),
            ('payment_status', 'Varchar(20)', 'Not Null', 'Payment status'),
            ('report_url', 'Varchar(500)', 'Not Null', 'Lab report URL'),
            ('report_results', 'JSON', 'Not Null', 'Test results data'),
            ('abnormal_flags', 'JSON', 'Not Null', 'Abnormal result flags'),
            ('ordered_at', 'DateTime', 'Not Null', 'Order creation timestamp'),
        ]
    },
    {
        'name': 'ehr_images',
        'description': 'Stores medical images uploaded to patient EHR wallet.',
        'fields': [
            ('image_id', 'UUID', 'Primary Key, Not Null', 'Unique image identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('image_type', 'Varchar(30)', 'Not Null', 'Image type (X-Ray/MRI/CT/Ultrasound)'),
            ('image_url', 'Varchar(500)', 'Not Null', 'Image storage URL'),
            ('title', 'Varchar(200)', 'Not Null', 'Image title'),
            ('description', 'Text', 'Not Null', 'Image description'),
            ('hospital_name', 'Varchar(200)', 'Not Null', 'Hospital where scan was done'),
            ('scan_date', 'Date', 'Null', 'Date of scan'),
            ('uploaded_by', 'UUID', 'FK to login_credentials, Null', 'Uploader reference'),
            ('uploaded_at', 'DateTime', 'Not Null', 'Upload timestamp'),
        ]
    },
    {
        'name': 'doctor_registrations',
        'description': 'Stores registered doctor profile and professional details.',
        'fields': [
            ('doctor_id', 'UUID', 'Primary Key, Not Null', 'Unique doctor identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Associated hospital reference'),
            ('dept_id', 'UUID', 'FK to departments, Null', 'Department reference'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Doctor full name'),
            ('specialization', 'Varchar(100)', 'Not Null', 'Medical specialization'),
            ('license_no', 'Varchar(50)', 'Unique, Not Null', 'Medical license number'),
            ('experience_years', 'Integer', 'Not Null', 'Years of experience'),
            ('consultation_fee', 'Decimal', 'Not Null', 'Consultation fee amount'),
            ('profile_photo', 'Varchar(500)', 'Not Null', 'Profile photo URL'),
            ('is_online', 'Boolean', 'Not Null', 'Online availability status'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Account approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'doctor_slots',
        'description': 'Stores doctor consultation time slots for appointment booking.',
        'fields': [
            ('slot_id', 'UUID', 'Primary Key, Not Null', 'Unique slot identifier'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations', 'Doctor reference'),
            ('slot_date', 'Date', 'Not Null', 'Appointment date'),
            ('start_time', 'Time', 'Not Null', 'Slot start time'),
            ('end_time', 'Time', 'Not Null', 'Slot end time'),
            ('consult_type', 'Varchar(20)', 'Not Null', 'Consultation type (online/offline)'),
            ('is_booked', 'Boolean', 'Not Null', 'Booking status flag'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'consultations',
        'description': 'Stores patient-doctor consultation records including video call details.',
        'fields': [
            ('consultation_id', 'UUID', 'Primary Key, Not Null', 'Unique consultation identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations', 'Doctor reference'),
            ('slot_id', 'UUID', 'FK to doctor_slots, Null', 'Booked slot reference'),
            ('jitsi_room_id', 'Varchar(200)', 'Not Null', 'Jitsi Meet room identifier'),
            ('status', 'Varchar(20)', 'Not Null', 'Status (scheduled/active/completed)'),
            ('ai_suggestions', 'JSON', 'Not Null', 'AI diagnosis suggestions'),
            ('doctor_notes', 'Text', 'Not Null', 'Doctor consultation notes'),
            ('final_diagnosis', 'Text', 'Not Null', 'Final diagnosis text'),
            ('to_emergency', 'Boolean', 'Not Null', 'Emergency escalation flag'),
            ('payment_status', 'Varchar(20)', 'Not Null', 'Payment status'),
            ('started_at', 'DateTime', 'Null', 'Consultation start timestamp'),
            ('ended_at', 'DateTime', 'Null', 'Consultation end timestamp'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'prescriptions',
        'description': 'Stores doctor-issued prescriptions with medicine details.',
        'fields': [
            ('prescription_id', 'UUID', 'Primary Key, Not Null', 'Unique prescription identifier'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations', 'Prescribing doctor reference'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('consultation_id', 'UUID', 'FK to consultations, Null', 'Related consultation reference'),
            ('medicines', 'JSON', 'Not Null', 'List of prescribed medicines with dosage'),
            ('diagnosis', 'Text', 'Not Null', 'Diagnosis for prescription'),
            ('instructions', 'Text', 'Not Null', 'Usage instructions'),
            ('is_verified', 'Boolean', 'Not Null', 'Prescription verification status'),
            ('valid_until', 'Date', 'Null', 'Prescription validity date'),
            ('pdf_url', 'Varchar(500)', 'Not Null', 'Generated PDF URL'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'pharmacist_registrations',
        'description': 'Stores registered pharmacist and pharmacy details.',
        'fields': [
            ('pharmacist_id', 'UUID', 'Primary Key, Not Null', 'Unique pharmacist identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('pharmacy_name', 'Varchar(200)', 'Not Null', 'Pharmacy name'),
            ('license_no', 'Varchar(50)', 'Unique, Not Null', 'Pharmacy license number'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Pharmacist full name'),
            ('address', 'Text', 'Not Null', 'Pharmacy address'),
            ('latitude', 'Decimal', 'Null', 'GPS latitude coordinate'),
            ('longitude', 'Decimal', 'Null', 'GPS longitude coordinate'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Account approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'medicine_orders',
        'description': 'Stores patient medicine orders with prescription verification and delivery tracking.',
        'fields': [
            ('med_order_id', 'UUID', 'Primary Key, Not Null', 'Unique order identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('pharmacist_id', 'UUID', 'FK to pharmacist_registrations, Null', 'Pharmacist reference'),
            ('prescription_id', 'UUID', 'FK to prescriptions, Null', 'Related prescription reference'),
            ('medicines', 'JSON', 'Not Null', 'Ordered medicines list'),
            ('total_amount', 'Decimal', 'Not Null', 'Total order amount'),
            ('payment_status', 'Varchar(20)', 'Not Null', 'Payment status'),
            ('delivery_address', 'Text', 'Not Null', 'Delivery address'),
            ('order_status', 'Varchar(30)', 'Not Null', 'Order status'),
            ('prescription_verified', 'Boolean', 'Not Null', 'Prescription verification flag'),
            ('requires_prescription', 'Boolean', 'Not Null', 'Prescription requirement flag'),
            ('delivery_otp', 'Varchar(6)', 'Not Null', 'Delivery OTP code'),
            ('otp_expiry', 'DateTime', 'Null', 'OTP expiry timestamp'),
            ('otp_verified', 'Boolean', 'Not Null', 'OTP verification flag'),
            ('estimated_delivery_days', 'Integer', 'Not Null', 'Estimated delivery days'),
            ('dispatched_at', 'DateTime', 'Null', 'Dispatch timestamp'),
            ('delivered_at', 'DateTime', 'Null', 'Delivery timestamp'),
            ('status_history', 'JSON', 'Not Null', 'Order status history log'),
            ('ordered_at', 'DateTime', 'Not Null', 'Order creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'pharmacy_inventory',
        'description': 'Stores pharmacy medicine inventory with stock and expiry details.',
        'fields': [
            ('inventory_id', 'UUID', 'Primary Key, Not Null', 'Unique inventory identifier'),
            ('pharmacy_id', 'UUID', 'FK to pharmacist_registrations', 'Pharmacy reference'),
            ('medicine_name', 'Varchar(200)', 'Not Null', 'Medicine name'),
            ('generic_name', 'Varchar(200)', 'Not Null', 'Generic medicine name'),
            ('category', 'Varchar(20)', 'Not Null', 'Category (tablet/syrup/injection/cream)'),
            ('description', 'Text', 'Not Null', 'Medicine description'),
            ('price_per_unit', 'Decimal', 'Not Null', 'Price per unit'),
            ('unit', 'Varchar(20)', 'Not Null', 'Unit type'),
            ('stock_quantity', 'Integer', 'Not Null', 'Available stock count'),
            ('reorder_level', 'Integer', 'Not Null', 'Minimum reorder threshold'),
            ('requires_prescription', 'Boolean', 'Not Null', 'Prescription requirement flag'),
            ('manufacturer', 'Varchar(200)', 'Not Null', 'Manufacturer name'),
            ('expiry_date', 'Date', 'Null', 'Medicine expiry date'),
            ('is_available', 'Boolean', 'Not Null', 'Availability flag'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'lab_tech_registrations',
        'description': 'Stores registered laboratory technician profile details.',
        'fields': [
            ('lab_tech_id', 'UUID', 'Primary Key, Not Null', 'Unique lab tech identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Associated hospital reference'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Lab technician full name'),
            ('qualification', 'Varchar(100)', 'Not Null', 'Educational qualification'),
            ('specialization', 'Varchar(100)', 'Not Null', 'Area of specialization'),
            ('phone', 'Varchar(15)', 'Not Null', 'Contact phone number'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Account approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'lab_orders',
        'description': 'Stores lab test orders created by doctors during consultations.',
        'fields': [
            ('order_id', 'UUID', 'Primary Key, Not Null', 'Unique order identifier'),
            ('doctor_id', 'UUID', 'FK to doctor_registrations', 'Ordering doctor reference'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('lab_tech_id', 'UUID', 'FK to lab_tech_registrations, Null', 'Assigned lab tech reference'),
            ('tests_ordered', 'JSON', 'Not Null', 'List of tests ordered'),
            ('priority', 'Varchar(20)', 'Not Null', 'Priority level (STAT/URGENT/NORMAL)'),
            ('status', 'Varchar(20)', 'Not Null', 'Order status'),
            ('notes', 'Text', 'Not Null', 'Additional notes'),
            ('payment_status', 'Varchar(20)', 'Not Null', 'Payment status'),
            ('ordered_at', 'DateTime', 'Not Null', 'Order creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'lab_reports',
        'description': 'Stores completed lab test reports with results and AI analysis.',
        'fields': [
            ('report_id', 'UUID', 'Primary Key, Not Null', 'Unique report identifier'),
            ('order_id', 'UUID', 'FK to lab_orders', 'Related lab order reference'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('results', 'JSON', 'Not Null', 'Test results data'),
            ('report_file_url', 'Varchar(500)', 'Not Null', 'Report file URL'),
            ('abnormal_flags', 'JSON', 'Not Null', 'Abnormal result indicators'),
            ('ai_analysis', 'Text', 'Not Null', 'AI-generated analysis'),
            ('saved_to_ehr', 'Boolean', 'Not Null', 'EHR save status flag'),
            ('uploaded_at', 'DateTime', 'Not Null', 'Report upload timestamp'),
        ]
    },
    {
        'name': 'ambulance_driver_registrations',
        'description': 'Stores registered ambulance driver profile details.',
        'fields': [
            ('driver_id', 'UUID', 'Primary Key, Not Null', 'Unique driver identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Associated hospital reference'),
            ('full_name', 'Varchar(100)', 'Not Null', 'Driver full name'),
            ('license_no', 'Varchar(50)', 'Unique, Not Null', 'Driving license number'),
            ('phone', 'Varchar(15)', 'Not Null', 'Contact phone number'),
            ('is_available', 'Boolean', 'Not Null', 'Availability status flag'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Account approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'ambulances',
        'description': 'Stores ambulance vehicle details and real-time GPS location.',
        'fields': [
            ('ambulance_id', 'UUID', 'Primary Key, Not Null', 'Unique ambulance identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Associated hospital reference'),
            ('driver_id', 'UUID', 'FK to ambulance_driver_registrations, Null', 'Assigned driver reference'),
            ('vehicle_no', 'Varchar(20)', 'Unique, Not Null', 'Vehicle registration number'),
            ('ambulance_type', 'Varchar(20)', 'Not Null', 'Type (basic/advanced/ICU)'),
            ('equipment_list', 'JSON', 'Not Null', 'Onboard equipment list'),
            ('is_available', 'Boolean', 'Not Null', 'Availability status flag'),
            ('current_lat', 'Decimal(10,7)', 'Null', 'Current GPS latitude'),
            ('current_lng', 'Decimal(10,7)', 'Null', 'Current GPS longitude'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'emergency_requests',
        'description': 'Stores patient emergency SOS requests with location and severity.',
        'fields': [
            ('emergency_id', 'UUID', 'Primary Key, Not Null', 'Unique emergency identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('triage_id', 'UUID', 'FK to triage_sessions, Null', 'Related triage session'),
            ('patient_lat', 'Decimal(10,7)', 'Not Null', 'Patient GPS latitude'),
            ('patient_lng', 'Decimal(10,7)', 'Not Null', 'Patient GPS longitude'),
            ('severity', 'Varchar(20)', 'Not Null', 'Severity level (low/moderate/high/critical)'),
            ('status', 'Varchar(20)', 'Not Null', 'Status (pending/dispatched/completed)'),
            ('assigned_hospital_id', 'UUID', 'FK to hospital_registrations, Null', 'Nearest hospital reference'),
            ('assigned_bed_id', 'UUID', 'FK to beds, Null', 'Assigned bed reference'),
            ('created_at', 'DateTime', 'Not Null', 'Emergency creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'ambulance_dispatch',
        'description': 'Stores ambulance dispatch records linking emergencies to ambulances.',
        'fields': [
            ('dispatch_id', 'UUID', 'Primary Key, Not Null', 'Unique dispatch identifier'),
            ('emergency_id', 'UUID', 'FK to emergency_requests', 'Emergency reference'),
            ('ambulance_id', 'UUID', 'FK to ambulances', 'Assigned ambulance reference'),
            ('dispatch_status', 'Varchar(30)', 'Not Null', 'Status (dispatched/en_route/arrived/completed)'),
            ('eta_minutes', 'Integer', 'Null', 'Estimated arrival time in minutes'),
            ('route_data', 'JSON', 'Not Null', 'Route waypoints data'),
            ('dispatched_at', 'DateTime', 'Not Null', 'Dispatch timestamp'),
            ('arrived_at', 'DateTime', 'Null', 'Arrival timestamp'),
            ('completed_at', 'DateTime', 'Null', 'Trip completion timestamp'),
        ]
    },
    {
        'name': 'vendor_registrations',
        'description': 'Stores registered medical equipment vendor details.',
        'fields': [
            ('vendor_id', 'UUID', 'Primary Key, Not Null', 'Unique vendor identifier'),
            ('login_id', 'UUID', 'Unique, FK to login_credentials', 'Reference to login credentials'),
            ('company_name', 'Varchar(200)', 'Not Null', 'Company name'),
            ('tax_id', 'Varchar(50)', 'Unique, Not Null', 'Tax identification number'),
            ('contact_name', 'Varchar(100)', 'Not Null', 'Primary contact person name'),
            ('phone', 'Varchar(15)', 'Not Null', 'Contact phone number'),
            ('business_license_url', 'Varchar(500)', 'Not Null', 'Business license document URL'),
            ('certifications', 'JSON', 'Not Null', 'List of certifications'),
            ('approval_status', 'Varchar(20)', 'Not Null', 'Account approval status'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Record last update timestamp'),
        ]
    },
    {
        'name': 'equipment_catalog',
        'description': 'Stores medical equipment listed by vendors for hospital purchase.',
        'fields': [
            ('product_id', 'UUID', 'Primary Key, Not Null', 'Unique product identifier'),
            ('vendor_id', 'UUID', 'FK to vendor_registrations', 'Vendor reference'),
            ('product_name', 'Varchar(200)', 'Not Null', 'Equipment product name'),
            ('category', 'Varchar(50)', 'Not Null', 'Equipment category'),
            ('specifications', 'JSON', 'Not Null', 'Technical specifications'),
            ('price', 'Decimal', 'Not Null', 'Product price'),
            ('stock_qty', 'Integer', 'Not Null', 'Available stock quantity'),
            ('image_url', 'Varchar(500)', 'Not Null', 'Product image URL'),
            ('listed_at', 'DateTime', 'Not Null', 'Listing creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'equipment_orders',
        'description': 'Stores hospital equipment orders with payment and OTP delivery tracking.',
        'fields': [
            ('eq_order_id', 'UUID', 'Primary Key, Not Null', 'Unique order identifier'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Ordering hospital reference'),
            ('vendor_id', 'UUID', 'FK to vendor_registrations', 'Vendor reference'),
            ('product_id', 'UUID', 'FK to equipment_catalog', 'Ordered product reference'),
            ('quantity', 'Integer', 'Not Null', 'Ordered quantity'),
            ('total_price', 'Decimal', 'Not Null', 'Total order price'),
            ('order_status', 'Varchar(30)', 'Not Null', 'Order status'),
            ('payment_status', 'Varchar(20)', 'Not Null', 'Payment status'),
            ('delivery_otp', 'Varchar(6)', 'Not Null', 'Delivery OTP code'),
            ('otp_expiry', 'DateTime', 'Null', 'OTP expiry timestamp'),
            ('otp_verified', 'Boolean', 'Not Null', 'OTP verification flag'),
            ('estimated_delivery_days', 'Integer', 'Not Null', 'Estimated delivery days'),
            ('dispatched_at', 'DateTime', 'Null', 'Dispatch timestamp'),
            ('delivered_at', 'DateTime', 'Null', 'Delivery timestamp'),
            ('status_history', 'JSON', 'Not Null', 'Order status history log'),
            ('ordered_at', 'DateTime', 'Not Null', 'Order creation timestamp'),
            ('updated_at', 'DateTime', 'Not Null', 'Last update timestamp'),
        ]
    },
    {
        'name': 'triage_sessions',
        'description': 'Stores AI-powered triage session results for patient symptom checking.',
        'fields': [
            ('triage_id', 'UUID', 'Primary Key, Not Null', 'Unique triage identifier'),
            ('patient_id', 'UUID', 'FK to patient_registrations', 'Patient reference'),
            ('symptoms_input', 'JSON', 'Not Null', 'Patient reported symptoms'),
            ('predicted_diseases', 'JSON', 'Not Null', 'AI predicted disease list with confidence'),
            ('confidence_score', 'Decimal', 'Null', 'Overall prediction confidence score'),
            ('severity', 'Varchar(20)', 'Not Null', 'Assessed severity level'),
            ('model_version', 'Varchar(50)', 'Not Null', 'AI model version used'),
            ('emergency_triggered', 'Boolean', 'Not Null', 'Emergency SOS trigger flag'),
            ('recommendation', 'Text', 'Not Null', 'AI recommendation text'),
            ('created_at', 'DateTime', 'Not Null', 'Session creation timestamp'),
        ]
    },
    {
        'name': 'fl_global_models',
        'description': 'Stores federated learning global model versions and performance metrics.',
        'fields': [
            ('model_id', 'UUID', 'Primary Key, Not Null', 'Unique model identifier'),
            ('version', 'Varchar(50)', 'Unique, Not Null', 'Model version string'),
            ('weights_file_url', 'Varchar(500)', 'Not Null', 'Model weights file URL'),
            ('accuracy', 'Decimal', 'Null', 'Model accuracy percentage'),
            ('hospitals_count', 'Integer', 'Not Null', 'Number of participating hospitals'),
            ('aggregation_algo', 'Varchar(50)', 'Not Null', 'Aggregation algorithm (FedAvg)'),
            ('is_active', 'Boolean', 'Not Null', 'Active model flag'),
            ('privacy_epsilon', 'Decimal', 'Null', 'Differential privacy epsilon value'),
            ('created_at', 'DateTime', 'Not Null', 'Model creation timestamp'),
        ]
    },
    {
        'name': 'fl_rounds',
        'description': 'Stores federated learning training round details and status.',
        'fields': [
            ('round_id', 'UUID', 'Primary Key, Not Null', 'Unique round identifier'),
            ('model_id', 'UUID', 'FK to fl_global_models', 'Parent model reference'),
            ('round_number', 'Integer', 'Not Null', 'Sequential round number'),
            ('status', 'Varchar(20)', 'Not Null', 'Round status (pending/active/completed)'),
            ('hospitals_invited', 'Integer', 'Not Null', 'Number of invited hospitals'),
            ('hospitals_completed', 'Integer', 'Not Null', 'Number of completed submissions'),
            ('global_loss', 'Decimal', 'Null', 'Global model loss value'),
            ('started_at', 'DateTime', 'Null', 'Round start timestamp'),
            ('completed_at', 'DateTime', 'Null', 'Round completion timestamp'),
            ('round_deadline', 'DateTime', 'Null', 'Submission deadline'),
            ('min_hospitals_threshold', 'Integer', 'Not Null', 'Minimum required hospitals'),
            ('auto_aggregated', 'Boolean', 'Not Null', 'Auto aggregation flag'),
            ('reminder_sent', 'Boolean', 'Not Null', 'Reminder notification sent flag'),
            ('created_at', 'DateTime', 'Not Null', 'Record creation timestamp'),
        ]
    },
    {
        'name': 'fl_hospital_weights',
        'description': 'Stores individual hospital model weight submissions for federated learning.',
        'fields': [
            ('weight_id', 'UUID', 'Primary Key, Not Null', 'Unique weight identifier'),
            ('round_id', 'UUID', 'FK to fl_rounds', 'FL round reference'),
            ('hospital_id', 'UUID', 'FK to hospital_registrations', 'Submitting hospital reference'),
            ('weights_file_url', 'Varchar(500)', 'Not Null', 'Weight file URL'),
            ('local_accuracy', 'Decimal', 'Null', 'Local model accuracy'),
            ('local_loss', 'Decimal', 'Null', 'Local model loss'),
            ('training_samples', 'Integer', 'Not Null', 'Number of training samples used'),
            ('noise_added', 'Boolean', 'Not Null', 'Differential privacy noise flag'),
            ('submitted_at', 'DateTime', 'Not Null', 'Submission timestamp'),
        ]
    },
    {
        'name': 'epidemic_trends',
        'description': 'Stores epidemic detection and disease trend monitoring data.',
        'fields': [
            ('trend_id', 'UUID', 'Primary Key, Not Null', 'Unique trend identifier'),
            ('disease_name', 'Varchar(100)', 'Not Null', 'Disease name'),
            ('region', 'Varchar(100)', 'Not Null', 'Affected region'),
            ('case_count', 'Integer', 'Not Null', 'Number of reported cases'),
            ('spike_detected', 'Boolean', 'Not Null', 'Epidemic spike detection flag'),
            ('heatmap_data', 'JSON', 'Not Null', 'Geographic heatmap data'),
            ('alert_level', 'Varchar(20)', 'Not Null', 'Alert level (low/medium/high/critical)'),
            ('recorded_date', 'Date', 'Not Null', 'Data recording date'),
        ]
    },
]


for table in tables_data:
    add_table_to_doc(
        doc,
        table['name'],
        table['description'],
        table['fields']
    )


output_path = 'federcare_database_tables.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
print(f'Total tables: {len(tables_data)}')
