"""
Generate Word document explaining each FederCare database table
in simple plain English from Table 1 to Table 41.
Output: federcare_table_explanations.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DARK_BLUE = RGBColor(26, 60, 110)


doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Header
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run('FederCare: AI Health Network')
hr.italic = True
hr.font.size = Pt(10)
hr.font.color.rgb = DARK_BLUE

# Footer
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
fr = fp.add_run('Mar Thoma Institute of Information Technology, Ayur')
fr.italic = True
fr.bold = True
fr.font.size = Pt(9)
fr.font.color.rgb = DARK_BLUE

fp.add_run('\t\t\t\t\t\t')

fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

prun = fp.add_run()
prun.element.append(fldChar1)
prun.element.append(instrText)
prun.element.append(fldChar2)
prun.font.size = Pt(9)
prun.bold = True


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6E')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table_block(doc, number, name, what, stores, connected):
    # Heading
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    hrun = heading.add_run(f'Table {number}: {name}')
    hrun.bold = True
    hrun.font.size = Pt(12)
    hrun.font.color.rgb = DARK_BLUE

    # What is this table?
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(4)
    lbl = p1.add_run('What is this table? ')
    lbl.bold = True
    lbl.font.size = Pt(10)
    txt = p1.add_run(what)
    txt.font.size = Pt(10)

    # What does it store?
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    lbl2 = p2.add_run('What does it store?')
    lbl2.bold = True
    lbl2.font.size = Pt(10)

    for item in stores:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Cm(1.0)
        bp.paragraph_format.space_after = Pt(1)
        brun = bp.add_run(item)
        brun.font.size = Pt(10)

    # How is it connected?
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(6)
    lbl3 = p3.add_run('How is it connected? ')
    lbl3.bold = True
    lbl3.font.size = Pt(10)
    txt3 = p3.add_run(connected)
    txt3.font.size = Pt(10)

    # Horizontal line
    line_para = doc.add_paragraph()
    add_horizontal_line(line_para)


# ===== Title page =====
doc.add_paragraph('\n\n')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('Database Tables — Simple Explanation Guide')
tr.bold = True
tr.font.size = Pt(20)
tr.font.color.rgb = DARK_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run('FederCare: AI Health Network')
sr.bold = True
sr.font.size = Pt(14)

doc.add_paragraph()

dept = doc.add_paragraph()
dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dept.add_run('Department of Computer Applications')
dr.font.size = Pt(12)

college = doc.add_paragraph()
college.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = college.add_run('Mar Thoma Institute of Information Technology, Ayur')
cr.font.size = Pt(12)

year = doc.add_paragraph()
year.alignment = WD_ALIGN_PARAGRAPH.CENTER
yr = year.add_run('2025 - 2026')
yr.font.size = Pt(12)

doc.add_page_break()


tables = [
    {
        'name': 'login_credentials',
        'what': 'This is the master authentication table that stores login details for every user in the FederCare system regardless of their role.',
        'stores': [
            'email — unique email used to login',
            'password_hash — encrypted password',
            'role — determines what the user can do (super_admin, hospital_admin, doctor, patient, pharmacist, lab_tech, driver, vendor)',
            'is_approved — admin must approve account before user can login',
        ],
        'connected': 'Every other user profile table (doctor, patient, hospital etc.) has a One-to-One link to this table through login_id. Think of this as the front door that all 8 types of users enter through.'
    },
    {
        'name': 'super_admin',
        'what': 'Stores the profile of the system super administrator who manages the entire FederCare platform.',
        'stores': [
            'full_name — administrator name',
            'phone — contact number',
            'profile_photo — photo URL',
        ],
        'connected': 'Links to login_credentials through login_id. There is only one super admin in the system.'
    },
    {
        'name': 'role_permissions',
        'what': 'Defines what each user role is allowed to do in each module of the system.',
        'stores': [
            'role — which user type (doctor, patient etc.)',
            'module — which part of system (EHR, FL etc.)',
            'can_read — can view data',
            'can_write — can add or edit data',
            'can_delete — can remove data',
        ],
        'connected': 'Standalone table. Used by the system to check permissions before every action.'
    },
    {
        'name': 'login_sessions',
        'what': 'Tracks every active login session to keep the system secure.',
        'stores': [
            'jwt_token_hash — encrypted session token',
            'device_info — which device logged in',
            'ip_address — where login came from',
            'expires_at — when session ends',
        ],
        'connected': 'Links to login_credentials. Each time a user logs in a new session record is created. When they logout or token expires the session ends.'
    },
    {
        'name': 'audit_logs',
        'what': 'Records every important action performed in the system for security and compliance tracking.',
        'stores': [
            'action — what was done',
            'module — which part of system',
            'old_value — data before change',
            'new_value — data after change',
            'logged_at — exactly when it happened',
        ],
        'connected': 'Links to login_credentials to record who did what. Like a CCTV for the database — every action is recorded.'
    },
    {
        'name': 'notifications',
        'what': 'Stores all notifications sent to users across the system for real-time alerts and updates.',
        'stores': [
            'title — notification heading',
            'message — notification content',
            'notif_type — order/emergency/alert/payment',
            'is_read — whether user has seen it',
        ],
        'connected': 'Links to login_credentials. Every notification goes to a specific user. When a patient pays for medicine the pharmacist gets a notification. When ambulance dispatches the patient gets a notification.'
    },
    {
        'name': 'hospital_registrations',
        'what': 'Stores complete profile and registration details of each hospital that joins the FederCare network.',
        'stores': [
            'hospital_name — official name',
            'registration_no — government license',
            'latitude/longitude — GPS location for finding nearest hospital in emergency',
            'telemedicine_enabled — supports video consultations',
            'approval_status — super admin must approve before hospital can use system',
        ],
        'connected': 'This is the most connected table after login_credentials. Doctors, lab techs, drivers, beds, departments, inventory all link to hospital_id. Think of hospital as the owner of all its resources.'
    },
    {
        'name': 'departments',
        'what': 'Stores the medical departments inside each hospital such as Cardiology, Orthopedics, Neurology etc.',
        'stores': [
            'dept_name — department name',
            'description — what the department does',
        ],
        'connected': 'Links to hospital_registrations. One hospital has many departments. Doctors are assigned to departments through dept_id in doctor_registrations.'
    },
    {
        'name': 'beds',
        'what': 'Tracks every hospital bed and whether it is available, occupied or reserved.',
        'stores': [
            'bed_type — general, ICU, or emergency',
            'ward_name — which ward the bed is in',
            'status — available/occupied/reserved',
            'reserved_for — links to patient who is using the bed',
        ],
        'connected': 'Links to hospital and patient. When an emergency patient arrives the system reserves a bed by updating reserved_for with the patient_id.'
    },
    {
        'name': 'hospital_inventory',
        'what': 'Tracks all medical equipment and supplies available in each hospital.',
        'stores': [
            'item_name — medicine or equipment name',
            'quantity — how many available',
            'reorder_level — alert threshold for low stock',
            'maintenance_due — when equipment needs servicing',
        ],
        'connected': 'Links to hospital_registrations. Hospital admin manages their own inventory. When quantity drops below reorder_level a low stock alert is sent.'
    },
    {
        'name': 'hospital_patients',
        'what': 'Stores patient records entered directly by hospital staff. These records are used as training data for the Federated Learning AI model.',
        'stores': [
            'full_name, age, gender, blood_group',
            'symptoms — list of symptoms in JSON',
            'diagnosis — what disease was diagnosed',
            'visit_date — when patient came',
        ],
        'connected': 'Links to hospital_registrations. Each hospital has their own set of patient records. The FL engine uses these records to train local models without sharing data with other hospitals. This is the key privacy feature of FederCare.'
    },
    {
        'name': 'patient_registrations',
        'what': 'Stores complete health profile of every patient registered in FederCare.',
        'stores': [
            'dob, gender, blood_group',
            'height_cm, weight_kg, bmi',
            'emergency_contact — who to call in an emergency',
            'qr_code_url — unique QR for EHR access',
            'lifestyle_data — exercise, diet habits',
        ],
        'connected': 'Most connected patient table. Links to consultations, prescriptions, lab orders, medicine orders, EHR records, emergency requests and more. Everything a patient does creates a record linked to their patient_id.'
    },
    {
        'name': 'ehr_records',
        'what': 'Stores all Electronic Health Records for each patient including diagnoses, test results and medical history.',
        'stores': [
            'record_type — prescription/lab/diagnosis',
            'title — record heading',
            'content — full record details',
            'file_url — attached document',
            'is_sensitive — marks sensitive records',
        ],
        'connected': 'Links to patient_registrations. Doctors and lab techs add records here. Patients can view their own records. Access to records requires QR consent.'
    },
    {
        'name': 'allergies',
        'what': 'Stores known allergies for each patient to prevent dangerous prescriptions.',
        'stores': [
            'allergen — what the patient is allergic to',
            'reaction — what happens when exposed',
            'severity — mild, moderate or severe',
            'noted_by — which doctor recorded it',
        ],
        'connected': 'Links to patient and doctor. When doctor prescribes medicines the system checks this table to warn about allergy conflicts.'
    },
    {
        'name': 'ehr_consent_log',
        'what': 'Tracks every time a doctor accesses a patient EHR to ensure privacy and consent compliance.',
        'stores': [
            'accessed_by — which doctor accessed',
            'access_type — view or download',
            'consent_given — patient gave permission',
            'expires_at — when the permission ends',
        ],
        'connected': 'Links to patient and login. Patient shows QR code to doctor. Doctor scans it. This creates a consent record valid for 30 minutes. After expiry doctor cannot access without new consent.'
    },
    {
        'name': 'risk_assessments',
        'what': 'Stores AI-generated health risk scores for each patient based on their health data.',
        'stores': [
            'diabetes_risk — percentage chance',
            'heart_risk — percentage chance',
            'hypertension_risk — percentage chance',
            'risk_level — overall low/medium/high',
            'recommendations — what patient should do',
        ],
        'connected': 'Links to patient_registrations. The AI model analyzes patient BMI, lifestyle data, and symptoms to calculate risk scores. High risk patients receive automatic alerts.'
    },
    {
        'name': 'patient_complaints',
        'what': 'Stores complaints filed by patients about doctors, hospitals or vendors.',
        'stores': [
            'complaint_type — doctor/vendor/hospital',
            'subject — complaint heading',
            'description — full complaint',
            'status — pending/reviewed/resolved',
            'admin_response — reply from admin',
        ],
        'connected': 'Links to patient, doctor, hospital and vendor. Doctor complaints go to hospital admin. Vendor complaints go to super admin. Both can respond through the system.'
    },
    {
        'name': 'lab_test_orders',
        'what': 'Stores lab test bookings made directly by patients through the patient portal.',
        'stores': [
            'tests — list of tests ordered',
            'total_fee — payment amount',
            'appointment_date/time — when to come',
            'report_url — link to completed report',
            'abnormal_flags — which results are outside normal range',
        ],
        'connected': 'Links to patient, hospital and doctor. Patient pays through Razorpay. Lab tech uploads results. Patient downloads report from their portal.'
    },
    {
        'name': 'ehr_images',
        'what': 'Stores medical images uploaded by patients to their EHR wallet such as X-Rays, MRI scans and CT scans.',
        'stores': [
            'image_type — X-Ray/MRI/CT/Ultrasound',
            'image_url — where file is stored',
            'hospital_name — where scan was done',
            'scan_date — when image was taken',
        ],
        'connected': 'Links to patient. Patients upload their own scans. Doctors can view these images during consultation if patient grants consent. The chest X-Ray AI model analyzes uploaded X-Ray images for pneumonia detection.'
    },
    {
        'name': 'doctor_registrations',
        'what': 'Stores complete professional profile of every doctor in FederCare.',
        'stores': [
            'specialization — area of expertise',
            'license_no — government license',
            'experience_years — years practicing',
            'consultation_fee — charge per session',
            'is_online — available for video call',
        ],
        'connected': 'Links to hospital and department. Doctor belongs to one hospital and one department. They have many consultation slots, consultations, prescriptions and lab orders linked to their doctor_id.'
    },
    {
        'name': 'doctor_slots',
        'what': 'Stores the available time slots that doctors create for patient appointment booking.',
        'stores': [
            'slot_date — date of appointment',
            'start_time/end_time — timing',
            'consult_type — online or offline',
            'is_booked — whether slot is taken',
        ],
        'connected': 'Links to doctor_registrations. When patient books a slot is_booked becomes True and a consultation record is created linking slot_id. One slot can only have one consultation.'
    },
    {
        'name': 'consultations',
        'what': 'Stores every patient-doctor consultation including video call details and AI diagnosis suggestions.',
        'stores': [
            'jitsi_room_id — unique video call link',
            'status — scheduled/active/completed',
            'ai_suggestions — AI predicted diseases shown during consultation',
            'doctor_notes — what doctor observed',
            'final_diagnosis — confirmed diagnosis',
            'payment_status — whether patient paid',
        ],
        'connected': 'Links patient, doctor and slot. This is the central record for an entire consultation. When doctor prescribes medicines prescription links to consultation_id. When doctor orders lab tests lab_order links to consultation_id.'
    },
    {
        'name': 'prescriptions',
        'what': 'Stores all prescriptions issued by doctors to patients during or after consultations.',
        'stores': [
            'medicines — JSON list of medicines with dosage, frequency and duration',
            'diagnosis — reason for prescription',
            'instructions — how to take medicines',
            'valid_until — expiry date',
            'pdf_url — downloadable prescription PDF',
        ],
        'connected': 'Links doctor, patient and consultation. Generated automatically after consultation. Patient can download PDF. Pharmacist verifies prescription before dispensing controlled medicines.'
    },
    {
        'name': 'pharmacist_registrations',
        'what': 'Stores registered pharmacy and pharmacist details in the FederCare network.',
        'stores': [
            'pharmacy_name — name of pharmacy',
            'license_no — pharmacy license',
            'full_name — pharmacist name',
            'latitude/longitude — pharmacy location',
        ],
        'connected': 'Links to login_credentials. Each pharmacist has their own inventory of medicines. Patients can see all pharmacies and their medicines in the medicine ordering section.'
    },
    {
        'name': 'medicine_orders',
        'what': 'Stores every medicine order placed by patients with full tracking from order to delivery.',
        'stores': [
            'medicines — what was ordered',
            'total_amount — order value',
            'order_status — current stage',
            'prescription_verified — pharmacist approved the prescription',
            'delivery_otp — 6-digit code for confirming delivery',
            'status_history — full audit trail of every status change',
        ],
        'connected': 'Links patient, pharmacist and prescription. If medicines need prescription the order waits for pharmacist verification before patient can pay. After payment pharmacist dispatches. Patient confirms with OTP.'
    },
    {
        'name': 'pharmacy_inventory',
        'what': 'Stores all medicines available in each pharmacy with stock levels and expiry dates.',
        'stores': [
            'medicine_name/generic_name',
            'category — tablet/syrup/injection',
            'price_per_unit — cost',
            'stock_quantity — available count',
            'requires_prescription — needs Rx or not',
            'expiry_date — medicine expiry',
        ],
        'connected': 'Links to pharmacist_registrations. Patients browse this table to select medicines. When order delivered stock_quantity decreases automatically. Low stock triggers notification to pharmacist.'
    },
    {
        'name': 'lab_tech_registrations',
        'what': 'Stores profile of laboratory technicians who process and upload test results.',
        'stores': [
            'qualification — educational degree',
            'specialization — area of expertise',
            'phone — contact number',
        ],
        'connected': 'Links to hospital. Lab tech sees pending orders in their hospital, processes tests, uploads results and marks abnormal values.'
    },
    {
        'name': 'lab_orders',
        'what': 'Stores lab test orders created by doctors during consultation for patients.',
        'stores': [
            'tests_ordered — list of required tests',
            'priority — STAT/URGENT/NORMAL',
            'status — pending/processing/completed',
            'notes — special instructions',
        ],
        'connected': 'Links doctor, patient and lab tech. Doctor creates during consultation. Lab tech picks up and processes. Results saved in lab_reports table linked to this order.'
    },
    {
        'name': 'lab_reports',
        'what': 'Stores completed lab test reports with results and AI analysis.',
        'stores': [
            'results — JSON with test values',
            'report_file_url — PDF report link',
            'abnormal_flags — which values are outside normal range',
            'ai_analysis — AI interpretation',
            'saved_to_ehr — linked to patient EHR',
        ],
        'connected': 'Links to lab_order and patient. After lab tech uploads report patient gets notification. Patient can download PDF from their portal. Abnormal results shown in red.'
    },
    {
        'name': 'ambulance_driver_registrations',
        'what': 'Stores profile of ambulance drivers registered in the FederCare emergency response system.',
        'stores': [
            'full_name — driver name',
            'license_no — driving license',
            'phone — contact number',
            'is_available — ready for dispatch',
        ],
        'connected': 'Links to hospital. Each driver is assigned one ambulance through the ambulances table. When emergency dispatched driver gets real-time alert via WebSocket.'
    },
    {
        'name': 'ambulances',
        'what': 'Stores ambulance vehicle details and real-time GPS location for emergency dispatch.',
        'stores': [
            'vehicle_no — registration number',
            'ambulance_type — basic/advanced/ICU',
            'equipment_list — medical equipment available on board',
            'is_available — free for dispatch',
            'current_lat/current_lng — live GPS position updated by driver',
        ],
        'connected': 'Links to hospital and driver. When emergency triggered system checks this table to find nearest available ambulance using GPS coordinates. is_available becomes False when dispatched and True again after hospital acknowledges patient delivery.'
    },
    {
        'name': 'emergency_requests',
        'what': 'Stores every emergency SOS request triggered by patients with their GPS location and severity.',
        'stores': [
            'patient_lat/patient_lng — where patient is',
            'severity — low/moderate/high/critical',
            'status — pending/dispatched/completed',
            'assigned_hospital_id — nearest hospital',
            'assigned_bed_id — reserved bed',
        ],
        'connected': 'Links patient, hospital and bed. When patient triggers SOS this record is created. System finds nearest ambulance using patient coordinates. AmbulanceDispatch record is created linking this emergency to the selected ambulance.'
    },
    {
        'name': 'ambulance_dispatch',
        'what': 'Stores the dispatch record that links an emergency to an ambulance and tracks the entire trip status.',
        'stores': [
            'dispatch_status — dispatched/en_route/arrived/pending_acknowledgment/completed',
            'eta_minutes — estimated arrival time',
            'route_data — GPS route coordinates',
            'dispatched_at — when sent',
            'arrived_at — when reached patient',
            'completed_at — when delivered',
        ],
        'connected': 'Links emergency_requests and ambulances. This is the tracking record for the entire ambulance journey. Status updates here trigger real-time notifications to patient and hospital. Hospital admin acknowledges delivery which frees the ambulance.'
    },
    {
        'name': 'vendor_registrations',
        'what': 'Stores registered medical equipment vendor companies that supply hospitals with equipment through FederCare.',
        'stores': [
            'company_name — vendor company',
            'tax_id — business tax number',
            'contact_name — representative name',
            'business_license_url — license document',
            'certifications — quality certificates',
        ],
        'connected': 'Links to login_credentials. Vendors list their products in equipment_catalog. Hospitals browse and order. Super admin approves vendors before they can list products.'
    },
    {
        'name': 'equipment_catalog',
        'what': 'Stores all medical equipment products listed by vendors for hospital purchase.',
        'stores': [
            'product_name — equipment name',
            'category — type of equipment',
            'specifications — technical details',
            'price — cost per unit',
            'stock_qty — available quantity',
            'image_url — product photo',
        ],
        'connected': 'Links to vendor_registrations. Hospital admin browses this catalog to order equipment. When order placed equipment_orders record created.'
    },
    {
        'name': 'equipment_orders',
        'what': 'Stores all equipment orders placed by hospitals with full payment and OTP delivery tracking.',
        'stores': [
            'quantity — how many ordered',
            'total_price — order cost',
            'order_status — current stage',
            'payment_status — paid or pending',
            'delivery_otp — 6-digit verification code',
            'status_history — complete order trail',
            'estimated_delivery_days — delivery ETA',
        ],
        'connected': 'Links hospital, vendor and equipment_catalog. After payment vendor dispatches. OTP sent to hospital. Hospital confirms delivery with OTP. This prevents fake delivery confirmations.'
    },
    {
        'name': 'triage_sessions',
        'what': 'Stores results of AI-powered symptom checking sessions where patients enter symptoms and AI predicts possible diseases.',
        'stores': [
            'symptoms_input — what patient reported',
            'predicted_diseases — AI predictions with confidence percentages',
            'confidence_score — overall accuracy',
            'severity — how serious the condition',
            'emergency_triggered — did patient need immediate emergency SOS',
            'recommendation — what AI advises',
        ],
        'connected': 'Links to patient. If AI detects critical symptoms it triggers emergency automatically. Results used by doctors during consultation as AI suggestions in consultations table.'
    },
    {
        'name': 'fl_global_models',
        'what': 'Stores each version of the global AI model trained through Federated Learning across all hospitals.',
        'stores': [
            'version — model version number',
            'accuracy — how accurate the model is',
            'hospitals_count — how many hospitals contributed to training',
            'aggregation_algo — FedAvg algorithm used',
            'is_active — currently in use',
            'privacy_epsilon — differential privacy strength applied',
        ],
        'connected': 'Links to fl_rounds. Each time FedAvg runs a new global model version is created with improved accuracy. Active model is used for all AI predictions across the platform.'
    },
    {
        'name': 'fl_rounds',
        'what': 'Stores each round of federated learning training where hospitals submit their local model weights.',
        'stores': [
            'round_number — which training round',
            'status — pending/active/completed',
            'hospitals_invited — how many asked',
            'hospitals_completed — how many submitted',
            'round_deadline — submission cutoff time',
            'min_hospitals_threshold — minimum needed before aggregation runs',
            'auto_aggregated — did it run automatically',
        ],
        'connected': 'Links to fl_global_models. Super admin starts a round. Hospitals get notified. Each hospital trains locally and submits weights. When enough hospitals submit FedAvg runs and new global model is created.'
    },
    {
        'name': 'fl_hospital_weights',
        'what': 'Stores the individual model weights submitted by each hospital during a federated learning training round.',
        'stores': [
            'weights_file_url — where weights stored',
            'local_accuracy — hospital model accuracy',
            'local_loss — hospital model loss',
            'training_samples — patients used',
            'noise_added — differential privacy applied to protect patient data',
        ],
        'connected': 'Links to fl_rounds and hospital_registrations. Each hospital submits their local weights here. The FedAvg algorithm reads all weights from this table and averages them to create the global model. Noise is added to protect individual patient records.'
    },
    {
        'name': 'epidemic_trends',
        'what': 'Stores epidemic detection alerts based on disease patterns detected across hospitals in the network.',
        'stores': [
            'disease_name — which disease spreading',
            'region — affected geographic area',
            'case_count — number of reported cases',
            'spike_detected — unusual increase found',
            'heatmap_data — geographic spread data',
            'alert_level — low/medium/high/critical',
        ],
        'connected': 'Populated automatically by the system when hospital_patients records show unusual patterns of same disease in multiple hospitals. Super admin sees epidemic alerts on dashboard and can broadcast warnings to all hospital staff.'
    },
]


for idx, t in enumerate(tables, start=1):
    add_table_block(doc, idx, t['name'], t['what'], t['stores'], t['connected'])


output_path = 'federcare_table_explanations.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
print(f'Total tables explained: {len(tables)}')
